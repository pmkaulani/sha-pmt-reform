from docx import Document
import os

docs_dir = r"c:\\Users\\STD USER\\Desktop\\work\\SHA\\docs"

def append_to_doc(filename, heading, text_blocks):
    filepath = os.path.join(docs_dir, filename)
    if not os.path.exists(filepath):
        print(f"Skipping {filename}, not found.")
        return

    try:
        doc = Document(filepath)
        doc.add_heading(heading, level=1)
        for text in text_blocks:
            doc.add_paragraph(text)
        doc.save(filepath)
        print(f"Successfully updated {filename}")
    except Exception as e:
        print(f"Error updating {filename}: {e}")

addendum_heading = "Addendum: Data Sources & Statistical Grounding"
addendum_text = [
    "To ensure the SHA Means-Testing Algorithm is fully defensible, all parameters, constraints, and multipliers are strictly grounded in official government data and regional health economic baselines:",
    "1. Demographic & Household Statistics (KDHS): The system uses the Kenya Demographic and Health Survey (KDHS 2022) to set biological and statistical constraints. The national average household size is 3.9. The algorithm sets a maximum reasonable threshold of 12 dependents. Any claim exceeding this statistically improbable threshold triggers a 'Phantom Dependents' fraud flag.",
    "2. Geographic and ASAL Adjustments (CRA & NDMA): The algorithm maps exactly to the 23 counties officially recognized by the Commission on Revenue Allocation (CRA) and the National Drought Management Authority (NDMA). The 85% land value discount in arid areas and 50% in semi-arid areas is justified by their classification as communal/survival assets rather than liquid real estate.",
    "3. Revenue & Compliance Stress Tests (KNBS & ILO): The baseline population of 15.5 million non-salaried households is derived from KNBS (Kenya National Bureau of Statistics) 2024 estimates. Compliance scenarios are based on historical NHIF engagement (Pessimistic - 30%), ILO informal sector taxation studies (Conservative - 45%), and the achieved rates of the Rwanda Community-Based Health Insurance model (Target - 60%).",
    "4. Legal & Privacy Frameworks (DPA & Constitution): All automated fraud escalations enforce the Data Protection Act (DPA) §35 and §39. The system mathematically prohibits automated rejections based solely on API mismatches to ensure compliance with the Constitutional right to fair administrative action (Article 47)."
]

append_to_doc("Algorithm_Documentation.docx", addendum_heading, addendum_text)
append_to_doc("Pitch_Deck_Outline.docx", addendum_heading, addendum_text)
append_to_doc("EXECUTIVE_SUMMARY_ONE_PAGER.docx", addendum_heading, addendum_text)
append_to_doc("ASSUMPTIONS.docx", addendum_heading, addendum_text)
append_to_doc("LEGAL_ETHICAL_ARCHITECTURAL_EVALUATION.docx", addendum_heading, addendum_text)

print("Finished appending stats to docs.")
