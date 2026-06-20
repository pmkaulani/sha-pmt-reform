
from docx import Document
import os

docs_dir = r"c:\Users\STD USER\Desktop\work\SHA\docs"

def append_to_doc(filename, heading, text_blocks):
    filepath = os.path.join(docs_dir, filename)
    if not os.path.exists(filepath):
        print(f"Skipping {filename}, not found.")
        return

    try:
        doc = Document(filepath)
        doc.add_heading(heading, level=1)
        for text in text_blocks:
            doc.add_paragraph(text)
        doc.save(filepath)
        print(f"Successfully updated {filename}")
    except Exception as e:
        print(f"Error updating {filename}: {e}")

addendum_heading = "Addendum: Resiliency & Advanced Fraud Prevention"
addendum_text = [
    "To ensure the system is completely legally defensible and resilient against API failures or deliberate manipulation, the following safeguards have been integrated into SHA PMT v2.1:",
    "1. Triangulation Fallback Protocol (Offline Resiliency): If KRA, NTSA, or Safaricom APIs fail, the algorithm does not crash or default to punitive overcharging. It mathematically degrades to utilize physical, verifiable proxies (e.g., land, livestock, physical vehicle valuation) with standard depreciation applied.",
    "2. 'Digital Ghost' Detection & CHP Protocol: If an applicant deliberately hides their digital footprint (claims no M-Pesa, no KRA, no NTSA) and shows zero assets to appear indigent, they are classified as a 'Digital Ghost'. They receive a provisional subsidy but are flagged for mandatory physical verification by a Community Health Promoter (CHP) within 90 days. This closes the loophole of tech-savvy individuals feigning poverty.",
    "3. Multi-SIM Concealment Risk: If an applicant declares a low-balance phone number, but Safaricom registry (via ID cross-check) shows multiple active SIM cards, the system flags a 'Multi-SIM Concealment Risk' requiring balance aggregation across all registered MSISDNs.",
    "4. Unverified Chama Treasurer Claim: The 80% Fiduciary Exemption is heavily guarded. If an applicant claims they are a Chama Treasurer to discount their liquidity, the system cross-references Safaricom/SACCO registries. If no registered group is linked to their ID/Phone, it triggers a CRITICAL 'Unverified Fiduciary Claim' flag and denies the exemption.",
    "5. Geographic Adjustments: The algorithm strictly applies 'Arid and Semi-Arid Land (ASAL)' adjustments. For example, 100 acres of land in Marsabit (arid) receives an 85% discount in valuation compared to 100 acres in Nairobi, recognizing it as illiquid capital rather than cash-equivalent wealth. Car valuations also scale by type (Standard Old vs. Luxury)."
]

append_to_doc("Algorithm_Documentation.docx", addendum_heading, addendum_text)
append_to_doc("Pitch_Deck_Outline.docx", addendum_heading, addendum_text)
append_to_doc("EXECUTIVE_SUMMARY_ONE_PAGER.docx", addendum_heading, addendum_text)
append_to_doc("LEGAL_ETHICAL_ARCHITECTURAL_EVALUATION.docx", addendum_heading, addendum_text)

